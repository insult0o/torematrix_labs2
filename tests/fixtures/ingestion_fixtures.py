"""
Test fixtures for document ingestion testing.

Provides utilities for creating test documents in various formats,
generating realistic test data, and setting up test environments.
"""

from pathlib import Path
import tempfile
from typing import List, Tuple, Dict, Any
import random
import string
import json
import csv
import zipfile
from datetime import datetime, timedelta
import uuid

# Try to import document creation libraries
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def create_test_pdf(filename: str, content: str, num_pages: int = 1) -> Path:
    """Create a test PDF file with specified content."""
    if not HAS_REPORTLAB:
        # Create a simple text file with .pdf extension for testing
        path = Path(filename)
        path.write_text(f"PDF Content: {content}")
        return path
    
    path = Path(filename)
    c = canvas.Canvas(str(path), pagesize=letter)
    
    # Split content into pages
    words = content.split()
    words_per_page = max(1, len(words) // num_pages)
    
    for page in range(num_pages):
        start_idx = page * words_per_page
        end_idx = start_idx + words_per_page if page < num_pages - 1 else len(words)
        page_content = " ".join(words[start_idx:end_idx])
        
        # Add page content
        y_position = 750
        for line in page_content.split('\n'):
            c.drawString(50, y_position, line[:80])  # Limit line length
            y_position -= 20
            if y_position < 50:  # Start new page if needed
                break
        
        c.showPage()
    
    c.save()
    return path


def create_test_docx(filename: str, content: str) -> Path:
    """Create a test DOCX file with specified content."""
    if not HAS_DOCX:
        # Create a simple text file with .docx extension for testing
        path = Path(filename)
        path.write_text(f"DOCX Content: {content}")
        return path
    
    doc = Document()
    doc.add_heading("Test Document", 0)
    
    # Split content into paragraphs
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
    
    doc.save(filename)
    return Path(filename)


def create_test_txt(filename: str, content: str) -> Path:
    """Create a test text file."""
    path = Path(filename)
    path.write_text(content, encoding='utf-8')
    return path


def create_test_html(filename: str, content: str, title: str = "Test Document") -> Path:
    """Create a test HTML file."""
    html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <title>{title}</title>
    <meta charset="utf-8">
</head>
<body>
    <h1>{title}</h1>
    <div>
        {content.replace('\n', '<br>\n')}
    </div>
</body>
</html>
"""
    path = Path(filename)
    path.write_text(html_content, encoding='utf-8')
    return path


def create_test_json(filename: str, data: Dict[str, Any]) -> Path:
    """Create a test JSON file."""
    path = Path(filename)
    with open(path, "w", encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    return path


def create_test_csv(filename: str, rows: List[List[str]], headers: List[str] = None) -> Path:
    """Create a test CSV file."""
    path = Path(filename)
    with open(path, "w", newline="", encoding='utf-8') as f:
        writer = csv.writer(f)
        if headers:
            writer.writerow(headers)
        writer.writerows(rows)
    return path


def create_test_zip(filename: str, file_contents: Dict[str, str]) -> Path:
    """Create a test ZIP archive with multiple files."""
    path = Path(filename)
    with zipfile.ZipFile(path, 'w') as zf:
        for file_name, content in file_contents.items():
            zf.writestr(file_name, content)
    return path


def create_test_files(specs: List[Tuple[str, str]], temp_dir: Path = None) -> List[Path]:
    """
    Create multiple test files based on specifications.
    
    Args:
        specs: List of (filename, content) tuples
        temp_dir: Directory to create files in (uses temp if None)
    
    Returns:
        List of created file paths
    """
    if temp_dir is None:
        temp_dir = Path(tempfile.mkdtemp())
    
    files = []
    
    for filename, content in specs:
        filepath = temp_dir / filename
        
        if filename.endswith('.pdf'):
            create_test_pdf(str(filepath), content)
        elif filename.endswith('.docx'):
            create_test_docx(str(filepath), content)
        elif filename.endswith('.txt'):
            create_test_txt(str(filepath), content)
        elif filename.endswith('.html'):
            create_test_html(str(filepath), content)
        elif filename.endswith('.json'):
            # Assume content is JSON string or create from content
            try:
                data = json.loads(content)
            except json.JSONDecodeError:
                data = {"content": content, "test": True}
            create_test_json(str(filepath), data)
        elif filename.endswith('.csv'):
            # Create simple CSV with content as rows
            rows = [line.split(',') for line in content.split('\n') if line.strip()]
            create_test_csv(str(filepath), rows)
        elif filename.endswith('.zip'):
            # Create ZIP with single file containing content
            create_test_zip(str(filepath), {"content.txt": content})
        else:
            # Default to text file
            create_test_txt(str(filepath), content)
        
        files.append(filepath)
    
    return files


def create_large_test_file(size_mb: int, filename: str = None) -> Path:
    """Create a large test file of specified size in MB."""
    if filename is None:
        filename = f"large_test_{size_mb}mb.bin"
    
    path = Path(filename)
    chunk_size = 1024 * 1024  # 1MB chunks
    
    with open(path, 'wb') as f:
        for _ in range(size_mb):
            # Generate random binary data
            chunk = bytes(random.getrandbits(8) for _ in range(chunk_size))
            f.write(chunk)
    
    return path


def generate_random_content(words: int = 100, sentences_per_paragraph: int = 5) -> str:
    """Generate random text content with realistic structure."""
    
    # Common words for more realistic text
    common_words = [
        "the", "and", "of", "to", "a", "in", "is", "it", "you", "that",
        "he", "was", "for", "on", "are", "as", "with", "his", "they",
        "document", "processing", "system", "data", "file", "content",
        "analysis", "information", "text", "paragraph", "section",
        "chapter", "report", "study", "research", "findings", "results"
    ]
    
    # Generate words
    word_list = []
    for _ in range(words):
        if random.random() < 0.7:  # 70% chance of common word
            word_list.append(random.choice(common_words))
        else:  # 30% chance of random word
            length = random.randint(3, 10)
            word = ''.join(random.choices(string.ascii_lowercase, k=length))
            word_list.append(word)
    
    # Group into sentences
    sentences = []
    i = 0
    while i < len(word_list):
        sentence_length = random.randint(5, 15)
        sentence_words = word_list[i:i + sentence_length]
        if sentence_words:
            sentence = ' '.join(sentence_words).capitalize() + '.'
            sentences.append(sentence)
        i += sentence_length
    
    # Group into paragraphs
    paragraphs = []
    i = 0
    while i < len(sentences):
        paragraph_sentences = sentences[i:i + sentences_per_paragraph]
        paragraph = ' '.join(paragraph_sentences)
        paragraphs.append(paragraph)
        i += sentences_per_paragraph
    
    return '\n\n'.join(paragraphs)


def create_test_dataset(num_files: int = 20, temp_dir: Path = None) -> Dict[str, List[Path]]:
    """
    Create a comprehensive test dataset with various file types.
    
    Returns:
        Dictionary mapping file types to lists of created files
    """
    if temp_dir is None:
        temp_dir = Path(tempfile.mkdtemp())
    
    dataset = {
        'pdf': [],
        'docx': [],
        'txt': [],
        'html': [],
        'json': [],
        'csv': [],
        'zip': []
    }
    
    # Create files of each type
    file_types = ['pdf', 'docx', 'txt', 'html', 'json', 'csv', 'zip']
    files_per_type = max(1, num_files // len(file_types))
    
    for file_type in file_types:
        for i in range(files_per_type):
            filename = f"test_{file_type}_{i+1}.{file_type}"
            content = generate_random_content(random.randint(50, 200))
            
            if file_type == 'json':
                # Create JSON structure
                json_data = {
                    "id": str(uuid.uuid4()),
                    "title": f"Test Document {i+1}",
                    "content": content,
                    "created_at": datetime.utcnow().isoformat(),
                    "metadata": {
                        "author": f"Test Author {i+1}",
                        "category": random.choice(["report", "document", "analysis"]),
                        "tags": [f"tag{j}" for j in range(random.randint(1, 5))]
                    }
                }
                filepath = create_test_json(str(temp_dir / filename), json_data)
            
            elif file_type == 'csv':
                # Create CSV data
                headers = ["id", "name", "category", "value", "date"]
                rows = []
                for j in range(random.randint(10, 50)):
                    rows.append([
                        str(j+1),
                        f"Item {j+1}",
                        random.choice(["A", "B", "C"]),
                        str(random.randint(1, 1000)),
                        datetime.utcnow().strftime("%Y-%m-%d")
                    ])
                filepath = create_test_csv(str(temp_dir / filename), rows, headers)
            
            elif file_type == 'zip':
                # Create ZIP with multiple files
                zip_contents = {
                    "readme.txt": "This is a test ZIP file",
                    "data.txt": content,
                    "info.json": json.dumps({"test": True, "files": 3})
                }
                filepath = create_test_zip(str(temp_dir / filename), zip_contents)
            
            else:
                # Create other file types
                specs = [(filename, content)]
                filepaths = create_test_files(specs, temp_dir)
                filepath = filepaths[0]
            
            dataset[file_type].append(filepath)
    
    return dataset


def create_corrupted_files(temp_dir: Path = None) -> List[Path]:
    """Create files with various corruption patterns for testing error handling."""
    if temp_dir is None:
        temp_dir = Path(tempfile.mkdtemp())
    
    corrupted_files = []
    
    # Empty file
    empty_file = temp_dir / "empty.pdf"
    empty_file.touch()
    corrupted_files.append(empty_file)
    
    # File with wrong extension
    wrong_ext = temp_dir / "text_as_pdf.pdf"
    wrong_ext.write_text("This is actually a text file, not a PDF")
    corrupted_files.append(wrong_ext)
    
    # Very large file name
    long_name = "a" * 200 + ".txt"
    long_name_file = temp_dir / long_name
    long_name_file.write_text("File with very long name")
    corrupted_files.append(long_name_file)
    
    # Binary junk with document extension
    junk_pdf = temp_dir / "junk.pdf"
    with open(junk_pdf, 'wb') as f:
        f.write(b'\x00\x01\x02\x03' * 1000)  # Random binary data
    corrupted_files.append(junk_pdf)
    
    # File with special characters
    special_chars = temp_dir / "special_chars_测试_файл.txt"
    special_chars.write_text("File with special characters in name", encoding='utf-8')
    corrupted_files.append(special_chars)
    
    return corrupted_files


def cleanup_test_files(file_paths: List[Path]):
    """Clean up test files after testing."""
    for file_path in file_paths:
        try:
            if file_path.exists():
                if file_path.is_file():
                    file_path.unlink()
                elif file_path.is_dir():
                    import shutil
                    shutil.rmtree(file_path)
        except Exception as e:
            print(f"Warning: Could not clean up {file_path}: {e}")


class TestDataGenerator:
    """Generator for creating realistic test data on demand."""
    
    def __init__(self, temp_dir: Path = None):
        self.temp_dir = temp_dir or Path(tempfile.mkdtemp())
        self._created_files = []
    
    def create_document_batch(self, count: int, file_types: List[str] = None) -> List[Path]:
        """Create a batch of documents for testing."""
        if file_types is None:
            file_types = ['pdf', 'docx', 'txt', 'html', 'json']
        
        files = []
        for i in range(count):
            file_type = random.choice(file_types)
            filename = f"batch_doc_{i+1}.{file_type}"
            content = generate_random_content(random.randint(100, 500))
            
            specs = [(filename, content)]
            created = create_test_files(specs, self.temp_dir)
            files.extend(created)
            self._created_files.extend(created)
        
        return files
    
    def create_realistic_report(self, filename: str = None) -> Path:
        """Create a realistic-looking report document."""
        if filename is None:
            filename = f"report_{datetime.now().strftime('%Y%m%d')}.txt"
        
        # Generate structured report content
        content = f"""
QUARTERLY REPORT - {datetime.now().strftime('%Y Q%m')}

EXECUTIVE SUMMARY
{generate_random_content(80, 3)}

METHODOLOGY
{generate_random_content(120, 4)}

FINDINGS AND ANALYSIS
{generate_random_content(200, 5)}

Key Performance Indicators:
- Metric A: {random.randint(50, 150)}%
- Metric B: {random.randint(1000, 5000)} units
- Metric C: ${random.randint(10000, 100000):,}

RECOMMENDATIONS
{generate_random_content(100, 4)}

CONCLUSION
{generate_random_content(60, 3)}

Prepared by: Test Analytics Team
Date: {datetime.now().strftime('%Y-%m-%d')}
        """.strip()
        
        filepath = create_test_txt(str(self.temp_dir / filename), content)
        self._created_files.append(filepath)
        return filepath
    
    def cleanup(self):
        """Clean up all created files."""
        cleanup_test_files(self._created_files)
        self._created_files = []
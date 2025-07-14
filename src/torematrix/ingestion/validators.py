"""
File validation service for document ingestion.

Provides comprehensive validation including MIME type verification,
content validation, and security checks.
"""

from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import magic
import hashlib
import logging
from datetime import datetime

# Optional imports for specific file type validation
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import zipfile
    HAS_ZIPFILE = True
except ImportError:
    HAS_ZIPFILE = False

logger = logging.getLogger(__name__)


class FileValidator(ABC):
    """Base validator interface for all file validators."""
    
    @abstractmethod
    async def validate(self, file_path: Path) -> List[str]:
        """
        Validate file and return list of errors.
        
        Args:
            file_path: Path to the file to validate
            
        Returns:
            List of error messages (empty if valid)
        """
        pass
    
    @property
    @abstractmethod
    def name(self) -> str:
        """Validator name for logging."""
        pass


class MimeTypeValidator(FileValidator):
    """Validates that MIME type matches file extension."""
    
    def __init__(self):
        """Initialize MIME type validator."""
        self._magic = magic.Magic(mime=True)
        self._extension_map = {
            ".pdf": ["application/pdf"],
            ".docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
            ".doc": ["application/msword"],
            ".txt": ["text/plain"],
            ".html": ["text/html"],
            ".xml": ["text/xml", "application/xml"],
            ".json": ["application/json"],
            ".csv": ["text/csv", "application/csv"],
            ".xlsx": ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"],
            ".xls": ["application/vnd.ms-excel"],
            ".pptx": ["application/vnd.openxmlformats-officedocument.presentationml.presentation"],
            ".ppt": ["application/vnd.ms-powerpoint"],
            ".png": ["image/png"],
            ".jpg": ["image/jpeg"],
            ".jpeg": ["image/jpeg"],
            ".gif": ["image/gif"],
            ".bmp": ["image/bmp"],
            ".tiff": ["image/tiff"],
            ".zip": ["application/zip", "application/x-zip-compressed"],
            ".rar": ["application/x-rar-compressed", "application/vnd.rar"],
            ".7z": ["application/x-7z-compressed"],
            ".md": ["text/markdown", "text/x-markdown", "text/plain"],
            ".rst": ["text/x-rst", "text/plain"],
            ".rtf": ["application/rtf", "text/rtf"],
            ".odt": ["application/vnd.oasis.opendocument.text"],
            ".epub": ["application/epub+zip"]
        }
    
    @property
    def name(self) -> str:
        return "MimeTypeValidator"
    
    async def validate(self, file_path: Path) -> List[str]:
        """Validate MIME type matches extension."""
        errors = []
        
        if not file_path.exists():
            errors.append(f"File not found: {file_path}")
            return errors
        
        ext = file_path.suffix.lower()
        
        try:
            detected_mime = self._magic.from_file(str(file_path))
            logger.debug(f"Detected MIME type: {detected_mime} for {file_path.name}")
            
            # Check if extension is in our map
            if ext in self._extension_map:
                expected_mimes = self._extension_map[ext]
                
                if detected_mime not in expected_mimes:
                    errors.append(
                        f"MIME type mismatch: expected {expected_mimes}, "
                        f"but detected {detected_mime}"
                    )
            else:
                # Unknown extension - log warning but don't fail
                logger.warning(f"Unknown file extension: {ext}")
                
        except Exception as e:
            errors.append(f"MIME type detection failed: {str(e)}")
            logger.error(f"MIME detection error for {file_path}: {e}")
        
        return errors


class FileSizeValidator(FileValidator):
    """Validates file size constraints."""
    
    def __init__(self, max_size: int = 100 * 1024 * 1024, min_size: int = 1):
        """
        Initialize file size validator.
        
        Args:
            max_size: Maximum file size in bytes (default 100MB)
            min_size: Minimum file size in bytes (default 1 byte)
        """
        self.max_size = max_size
        self.min_size = min_size
    
    @property
    def name(self) -> str:
        return "FileSizeValidator"
    
    async def validate(self, file_path: Path) -> List[str]:
        """Validate file size is within limits."""
        errors = []
        
        if not file_path.exists():
            errors.append(f"File not found: {file_path}")
            return errors
        
        try:
            size = file_path.stat().st_size
            
            if size < self.min_size:
                errors.append(f"File too small: {size} bytes (minimum: {self.min_size})")
            
            if size > self.max_size:
                errors.append(
                    f"File too large: {size:,} bytes "
                    f"(maximum: {self.max_size:,} bytes)"
                )
                
        except Exception as e:
            errors.append(f"Failed to check file size: {str(e)}")
            logger.error(f"Size check error for {file_path}: {e}")
        
        return errors


class ContentValidator(FileValidator):
    """Validates file content is readable and not corrupted."""
    
    @property
    def name(self) -> str:
        return "ContentValidator"
    
    async def validate(self, file_path: Path) -> List[str]:
        """Validate file content based on type."""
        errors = []
        
        if not file_path.exists():
            errors.append(f"File not found: {file_path}")
            return errors
        
        ext = file_path.suffix.lower()
        
        try:
            if ext == ".pdf":
                errors.extend(await self._validate_pdf(file_path))
            elif ext in [".docx", ".doc"]:
                errors.extend(await self._validate_word(file_path))
            elif ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp"]:
                errors.extend(await self._validate_image(file_path))
            elif ext in [".zip", ".rar", ".7z"]:
                errors.extend(await self._validate_archive(file_path))
            elif ext in [".txt", ".md", ".rst", ".csv"]:
                errors.extend(await self._validate_text(file_path))
            elif ext in [".json"]:
                errors.extend(await self._validate_json(file_path))
            elif ext in [".xml", ".html"]:
                errors.extend(await self._validate_xml(file_path))
                
        except Exception as e:
            errors.append(f"Content validation failed: {str(e)}")
            logger.error(f"Content validation error for {file_path}: {e}")
        
        return errors
    
    async def _validate_pdf(self, file_path: Path) -> List[str]:
        """Validate PDF file content."""
        errors = []
        
        if not HAS_PYPDF2:
            logger.warning("PyPDF2 not installed, skipping PDF content validation")
            return errors
        
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                
                # Check if encrypted
                if reader.is_encrypted:
                    errors.append("PDF is encrypted")
                
                # Check page count
                page_count = len(reader.pages)
                if page_count == 0:
                    errors.append("PDF has no pages")
                else:
                    logger.debug(f"PDF has {page_count} pages")
                
                # Try to read first page to verify it's not corrupted
                if page_count > 0:
                    try:
                        page = reader.pages[0]
                        text = page.extract_text()
                        if not text or not text.strip():
                            # Don't error on this - could be scanned PDF
                            logger.info("PDF appears to have no extractable text (may be scanned)")
                    except Exception as e:
                        errors.append(f"Cannot read PDF content: {str(e)}")
                        
        except PyPDF2.errors.PdfReadError as e:
            errors.append(f"Invalid or corrupted PDF: {str(e)}")
        except Exception as e:
            errors.append(f"PDF validation error: {str(e)}")
        
        return errors
    
    async def _validate_word(self, file_path: Path) -> List[str]:
        """Validate Word document content."""
        errors = []
        
        if not HAS_DOCX:
            logger.warning("python-docx not installed, skipping Word content validation")
            return errors
        
        if file_path.suffix.lower() == ".docx":
            try:
                doc = docx.Document(str(file_path))
                
                # Check if document has content
                if not doc.paragraphs and not doc.tables:
                    errors.append("Word document appears to be empty")
                    
            except Exception as e:
                errors.append(f"Invalid or corrupted Word document: {str(e)}")
        else:
            # .doc files require different handling
            logger.info("Legacy .doc format - content validation limited")
        
        return errors
    
    async def _validate_image(self, file_path: Path) -> List[str]:
        """Validate image file content."""
        errors = []
        
        if not HAS_PIL:
            logger.warning("PIL not installed, skipping image content validation")
            return errors
        
        try:
            with Image.open(file_path) as img:
                # Verify image can be read
                img.verify()
                
                # Get basic info
                logger.debug(f"Image format: {img.format}, size: {img.size}, mode: {img.mode}")
                
        except Exception as e:
            errors.append(f"Invalid or corrupted image: {str(e)}")
        
        return errors
    
    async def _validate_archive(self, file_path: Path) -> List[str]:
        """Validate archive file content."""
        errors = []
        
        ext = file_path.suffix.lower()
        
        if ext == ".zip" and HAS_ZIPFILE:
            try:
                with zipfile.ZipFile(file_path, 'r') as zf:
                    # Test archive integrity
                    bad_files = zf.testzip()
                    if bad_files:
                        errors.append(f"Corrupted files in ZIP: {bad_files}")
                    
                    # Check if empty
                    if not zf.namelist():
                        errors.append("ZIP archive is empty")
                        
            except zipfile.BadZipFile:
                errors.append("Invalid or corrupted ZIP file")
            except Exception as e:
                errors.append(f"ZIP validation error: {str(e)}")
        else:
            logger.info(f"Archive format {ext} validation not implemented")
        
        return errors
    
    async def _validate_text(self, file_path: Path) -> List[str]:
        """Validate text file content."""
        errors = []
        
        try:
            # Try to read as text with encoding detection
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read(1024)  # Read first 1KB
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                errors.append("Unable to decode text file with common encodings")
            elif not content.strip():
                errors.append("Text file appears to be empty")
                
        except Exception as e:
            errors.append(f"Text validation error: {str(e)}")
        
        return errors
    
    async def _validate_json(self, file_path: Path) -> List[str]:
        """Validate JSON file content."""
        errors = []
        
        try:
            import json
            with open(file_path, 'r') as f:
                json.load(f)
        except json.JSONDecodeError as e:
            errors.append(f"Invalid JSON: {str(e)}")
        except Exception as e:
            errors.append(f"JSON validation error: {str(e)}")
        
        return errors
    
    async def _validate_xml(self, file_path: Path) -> List[str]:
        """Validate XML/HTML file content."""
        errors = []
        
        try:
            import xml.etree.ElementTree as ET
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            if root is None:
                errors.append("XML file has no root element")
                
        except ET.ParseError as e:
            # For HTML, this might be expected
            if file_path.suffix.lower() == ".xml":
                errors.append(f"Invalid XML: {str(e)}")
            else:
                logger.debug(f"HTML parse warning (may be normal): {e}")
        except Exception as e:
            errors.append(f"XML validation error: {str(e)}")
        
        return errors


class SecurityValidator(FileValidator):
    """Validates files for security threats."""
    
    def __init__(self, check_macros: bool = True):
        """
        Initialize security validator.
        
        Args:
            check_macros: Whether to check for macros in Office documents
        """
        self.check_macros = check_macros
    
    @property
    def name(self) -> str:
        return "SecurityValidator"
    
    async def validate(self, file_path: Path) -> List[str]:
        """Perform security validation checks."""
        errors = []
        
        if not file_path.exists():
            errors.append(f"File not found: {file_path}")
            return errors
        
        # Check for suspicious file patterns
        suspicious_patterns = [
            ".exe", ".bat", ".cmd", ".com", ".scr", ".vbs", ".js",
            ".jar", ".app", ".dmg", ".pkg", ".deb", ".rpm"
        ]
        
        ext = file_path.suffix.lower()
        if ext in suspicious_patterns:
            errors.append(f"Potentially dangerous file type: {ext}")
        
        # Check for double extensions
        name_parts = file_path.name.split('.')
        if len(name_parts) > 2:
            # Check for dangerous double extensions
            if name_parts[-2].lower() in ["pdf", "doc", "docx", "txt", "jpg", "png"]:
                if name_parts[-1].lower() in ["exe", "bat", "scr", "com"]:
                    errors.append(f"Suspicious double extension: {file_path.name}")
        
        # Additional checks could be added here:
        # - Virus scanning integration
        # - Macro detection in Office files
        # - Script detection in PDFs
        
        return errors


class HashValidator(FileValidator):
    """Validates file integrity using hash verification."""
    
    def __init__(self, algorithm: str = "sha256"):
        """
        Initialize hash validator.
        
        Args:
            algorithm: Hash algorithm to use (sha256, md5, etc.)
        """
        self.algorithm = algorithm
    
    @property
    def name(self) -> str:
        return "HashValidator"
    
    async def validate(self, file_path: Path) -> List[str]:
        """Calculate file hash (doesn't validate, just computes)."""
        errors = []
        
        if not file_path.exists():
            errors.append(f"File not found: {file_path}")
            return errors
        
        try:
            hash_obj = hashlib.new(self.algorithm)
            
            with open(file_path, "rb") as f:
                while chunk := f.read(8192):
                    hash_obj.update(chunk)
            
            file_hash = hash_obj.hexdigest()
            logger.debug(f"File hash ({self.algorithm}): {file_hash}")
            
        except Exception as e:
            errors.append(f"Hash calculation failed: {str(e)}")
        
        return errors
    
    async def calculate_hash(self, file_path: Path) -> str:
        """Calculate and return file hash."""
        hash_obj = hashlib.new(self.algorithm)
        
        with open(file_path, "rb") as f:
            while chunk := f.read(8192):
                hash_obj.update(chunk)
        
        return hash_obj.hexdigest()


class CompositeValidator(FileValidator):
    """Runs multiple validators in sequence."""
    
    def __init__(self, validators: List[FileValidator]):
        """
        Initialize composite validator.
        
        Args:
            validators: List of validators to run
        """
        self.validators = validators
    
    @property
    def name(self) -> str:
        return "CompositeValidator"
    
    async def validate(self, file_path: Path) -> List[str]:
        """Run all validators and collect errors."""
        all_errors = []
        
        for validator in self.validators:
            try:
                logger.debug(f"Running {validator.name} on {file_path.name}")
                errors = await validator.validate(file_path)
                
                if errors:
                    # Prefix errors with validator name
                    prefixed_errors = [f"[{validator.name}] {error}" for error in errors]
                    all_errors.extend(prefixed_errors)
                    
            except Exception as e:
                error_msg = f"[{validator.name}] Validator failed: {str(e)}"
                all_errors.append(error_msg)
                logger.error(f"Validator {validator.name} failed: {e}")
        
        return all_errors


def create_default_validator(
    max_file_size: int = 100 * 1024 * 1024,
    check_security: bool = True
) -> CompositeValidator:
    """
    Create a default composite validator with standard checks.
    
    Args:
        max_file_size: Maximum file size in bytes
        check_security: Whether to include security checks
        
    Returns:
        Configured composite validator
    """
    validators = [
        FileSizeValidator(max_size=max_file_size),
        MimeTypeValidator(),
        ContentValidator(),
    ]
    
    if check_security:
        validators.append(SecurityValidator())
    
    return CompositeValidator(validators)